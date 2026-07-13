"""Build the controlled Word edition of the WAMOCON marketing handbook.

The Markdown handbook remains the review-friendly source. This script applies
the project reference-guide style, preserves headings/tables/lists, creates a
real Word table of contents field, and emits no environment secrets.
"""

from __future__ import annotations

import argparse
from pathlib import Path
import re

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


NAVY = "17324D"
TEAL = "167D86"
INK = "253746"
MUTED = "5F6F7A"
PALE = "EAF4F4"
PALE_BLUE = "EAF0F5"
PALE_RED = "FBECEC"
RULE = "CBD5DC"
WHITE = "FFFFFF"


def _set_cell_fill(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def _set_cell_margins(cell, *, top: int = 90, start: int = 100, bottom: int = 90, end: int = 100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def _keep_row_together(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def _set_paragraph_shading(paragraph, color: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    p_pr.append(shd)


def _set_keep(paragraph, *, with_next: bool = False) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    keep_lines = OxmlElement("w:keepLines")
    p_pr.append(keep_lines)
    if with_next:
        keep_next = OxmlElement("w:keepNext")
        p_pr.append(keep_next)


def _add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instruction, separate, text, end):
        run._r.append(element)


def _add_toc(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = ' TOC \\o "1-3" \\h \\z \\u '
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Inhaltsverzeichnis wird beim Öffnen aktualisiert."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instruction, separate, placeholder, end):
        run._r.append(element)


def _add_hyperlink(paragraph, text: str, target: str) -> None:
    relationship_id = paragraph.part.relate_to(
        target,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), TEAL)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.extend((color, underline))
    run.append(run_properties)
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


INLINE_PATTERN = re.compile(r"(\[([^\]]+)\]\(([^)]+)\)|\*\*([^*]+)\*\*|`([^`]+)`)")


def _inline(paragraph, text: str) -> None:
    cursor = 0
    for match in INLINE_PATTERN.finditer(text):
        if match.start() > cursor:
            paragraph.add_run(text[cursor : match.start()])
        token = match.group(0)
        if token.startswith("["):
            _add_hyperlink(paragraph, match.group(2), match.group(3))
        elif token.startswith("**"):
            run = paragraph.add_run(match.group(4))
            run.bold = True
        else:
            run = paragraph.add_run(match.group(5))
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor.from_string(NAVY)
        cursor = match.end()
    if cursor < len(text):
        paragraph.add_run(text[cursor:])


def _configure_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.2
    normal.paragraph_format.widow_control = True

    for name, size, color, before, after in (
        ("Title", 30, NAVY, 0, 12),
        ("Subtitle", 15, TEAL, 0, 12),
        ("Heading 1", 21, NAVY, 18, 9),
        ("Heading 2", 15, TEAL, 14, 7),
        ("Heading 3", 11.5, NAVY, 10, 5),
        ("Heading 4", 10.5, NAVY, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    if "Handbook Code" not in styles:
        code_style = styles.add_style("Handbook Code", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code_style = styles["Handbook Code"]
    code_style.font.name = "Consolas"
    code_style.font.size = Pt(8.5)
    code_style.font.color.rgb = RGBColor.from_string(INK)
    code_style.paragraph_format.left_indent = Inches(0.18)
    code_style.paragraph_format.right_indent = Inches(0.18)
    code_style.paragraph_format.space_before = Pt(3)
    code_style.paragraph_format.space_after = Pt(3)


def _configure_sections(document: Document) -> None:
    for section in document.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)
        section.header_distance = Inches(0.32)
        section.footer_distance = Inches(0.3)

        header = section.header
        paragraph = header.paragraphs[0]
        paragraph.text = "WAMOCON  ·  MARKETING-MASCHINE  ·  BETRIEBSHANDBUCH"
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = paragraph.runs[0]
        run.font.name = "Calibri"
        run.font.size = Pt(8)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(MUTED)

        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run("INTERN · KONTROLLIERTE KOPIE   |   ")
        run.font.name = "Calibri"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(MUTED)
        _add_page_field(paragraph)


def _cover(
    document: Document,
    *,
    metadata: dict[str, str],
    status_date: str,
    status_label: str,
) -> None:
    banner = document.add_table(rows=1, cols=1)
    banner.alignment = WD_TABLE_ALIGNMENT.CENTER
    banner.autofit = False
    banner.columns[0].width = Inches(6.8)
    cell = banner.cell(0, 0)
    _set_cell_fill(cell, NAVY)
    _set_cell_margins(cell, top=260, bottom=260, start=260, end=260)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("WAMOCON")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string(WHITE)

    document.add_paragraph("")
    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Marketing-Maschine")
    subtitle = document.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Betriebs- und Anwenderhandbuch")
    strapline = document.add_paragraph()
    strapline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = strapline.add_run("Einfach für Marketing · kontrolliert im Hintergrund")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string(MUTED)

    document.add_paragraph("")
    status = document.add_table(rows=1, cols=1)
    status.alignment = WD_TABLE_ALIGNMENT.CENTER
    status.autofit = False
    status.columns[0].width = Inches(6.2)
    cell = status.cell(0, 0)
    _set_cell_fill(cell, PALE_RED)
    _set_cell_margins(cell, top=180, bottom=180, start=220, end=220)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(f"BETRIEBSSTATUS · {status_date.upper()}\n")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(NAVY)
    run = paragraph.add_run(status_label.upper())
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0xA1, 0x23, 0x23)
    paragraph.add_run(
        "\nNeue Erstellung, Freigaben und externe Übergaben bleiben gesperrt, bis die dokumentierte Release-Abnahme vollständig grün ist."
    )

    document.add_paragraph("")
    values = (
        ("Version", metadata["Version"]),
        ("Stand", metadata["Stand"]),
        ("Status", metadata["Status"]),
        ("Dokumentverantwortung", metadata["Dokumentverantwortung"]),
        ("Technische Verantwortung", metadata["Technische Verantwortung"]),
        ("Gültiger Umfang", metadata["Gültiger Umfang"]),
    )
    meta = document.add_table(rows=len(values), cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row, (label, value) in zip(meta.rows, values):
        row.cells[0].text = label
        row.cells[1].text = value
        row.cells[0].paragraphs[0].runs[0].bold = True
        for cell in row.cells:
            _set_cell_margins(cell, top=70, bottom=70, start=90, end=90)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)
    document.add_paragraph("")
    notice = document.add_paragraph()
    notice.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = notice.add_run("Keine Zugangsdaten · keine privaten Belege · keine Sammelkonten")
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    document.add_page_break()


def _table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    columns = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=columns)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    compact_size = 8.2 if columns >= 5 else 9
    for row_index, values in enumerate(rows):
        row = table.rows[row_index]
        _keep_row_together(row)
        if row_index == 0:
            _repeat_table_header(row)
        for column_index in range(columns):
            cell = row.cells[column_index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_margins(cell)
            if row_index == 0:
                _set_cell_fill(cell, NAVY)
            elif row_index % 2 == 0:
                _set_cell_fill(cell, "F5F8FA")
            paragraph = cell.paragraphs[0]
            paragraph.clear()
            _inline(paragraph, values[column_index] if column_index < len(values) else "")
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                run.font.size = Pt(compact_size)
                if row_index == 0:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(WHITE)
    document.add_paragraph("").paragraph_format.space_after = Pt(1)


def _callout(document: Document, lines: list[str]) -> None:
    text = " ".join(line.strip()[1:].strip() for line in lines)
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    _set_cell_fill(cell, PALE_RED if "NICHT produktionsbereit" in text or "Aktueller Stopp" in text else PALE_BLUE)
    _set_cell_margins(cell, top=130, bottom=130, start=170, end=170)
    paragraph = cell.paragraphs[0]
    _inline(paragraph, text)
    paragraph.paragraph_format.space_after = Pt(0)
    _set_keep(paragraph)
    document.add_paragraph("").paragraph_format.space_after = Pt(1)


def _parse_markdown(document: Document, source: str) -> None:
    lines = source.splitlines()
    first_rule = next((index for index, line in enumerate(lines) if line.strip() == "---"), 0)
    lines = lines[first_rule + 1 :]
    index = 0
    in_code = False
    code_lines: list[str] = []
    while index < len(lines):
        raw = lines[index]
        stripped = raw.strip()
        if stripped.startswith("```"):
            if in_code:
                paragraph = document.add_paragraph(style="Handbook Code")
                paragraph.add_run("\n".join(code_lines))
                _set_paragraph_shading(paragraph, "F1F4F6")
                _set_keep(paragraph)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(raw)
            index += 1
            continue
        if not stripped:
            index += 1
            continue
        if stripped == "---":
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(6)
            p_pr = paragraph._p.get_or_add_pPr()
            borders = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "8")
            bottom.set(qn("w:color"), TEAL)
            borders.append(bottom)
            p_pr.append(borders)
            index += 1
            continue
        if stripped.startswith("|") and index + 1 < len(lines) and re.match(r"^\s*\|?\s*:?-+", lines[index + 1]):
            table_lines = [stripped]
            index += 2
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            rows = [[cell.strip() for cell in line.strip("|").split("|")] for line in table_lines]
            _table(document, rows)
            continue
        if stripped.startswith(">"):
            quote_lines = []
            while index < len(lines) and lines[index].strip().startswith(">"):
                quote_lines.append(lines[index].strip())
                index += 1
            _callout(document, quote_lines)
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading:
            level = min(4, len(heading.group(1)))
            paragraph = document.add_heading(level=level)
            _inline(paragraph, heading.group(2))
            _set_keep(paragraph, with_next=True)
            index += 1
            continue
        bullet = re.match(r"^(\s*)[-*]\s+(.+)$", raw)
        numbered = re.match(r"^(\s*)(\d+)\.\s+(.+)$", raw)
        if bullet or numbered:
            match = bullet if bullet else numbered
            assert match is not None
            paragraph = document.add_paragraph(style="List Bullet" if bullet else None)
            level = min(2, len(match.group(1).replace("\t", "    ")) // 2)
            paragraph.paragraph_format.left_indent = Inches(0.25 + level * 0.22)
            paragraph.paragraph_format.first_line_indent = Inches(-0.18)
            paragraph.paragraph_format.space_after = Pt(2)
            if numbered:
                marker = paragraph.add_run(f"{numbered.group(2)}. ")
                marker.bold = True
                _inline(paragraph, numbered.group(3))
            else:
                assert bullet is not None
                _inline(paragraph, bullet.group(2))
            _set_keep(paragraph)
            index += 1
            continue
        paragraph = document.add_paragraph()
        _inline(paragraph, stripped.rstrip())
        _set_keep(paragraph)
        index += 1


DOCUMENT_CONTROL_FIELDS = (
    "Version",
    "Stand",
    "Status",
    "Dokumentverantwortung",
    "Technische Verantwortung",
    "Gültiger Umfang",
)


def _document_metadata(source_text: str) -> tuple[dict[str, str], str, str]:
    metadata: dict[str, str] = {}
    for field in DOCUMENT_CONTROL_FIELDS:
        match = re.search(
            rf"^\|\s*{re.escape(field)}\s*\|\s*([^|]+?)\s*\|$",
            source_text,
            re.MULTILINE,
        )
        if match is None:
            raise ValueError(f"handbook source is missing document-control {field} row")
        metadata[field] = match.group(1).strip()

    version = metadata["Version"]
    if not re.fullmatch(r"\d+\.\d+", version):
        raise ValueError("handbook document-control version must use major.minor format")

    status_match = re.search(
        r"^>\s*\*\*Betriebsstatus am\s+(.+?):\s*(.+?)\.?\*\*\s*$",
        source_text,
        re.MULTILINE,
    )
    if status_match is None:
        raise ValueError("handbook source is missing the controlled Betriebsstatus callout")
    status_date = status_match.group(1).strip()
    status_label = status_match.group(2).strip().rstrip(".")
    return metadata, status_date, status_label


def build(source: Path, output: Path) -> None:
    source_text = source.read_text(encoding="utf-8")
    metadata, status_date, status_label = _document_metadata(source_text)

    document = Document()
    _configure_styles(document)
    _configure_sections(document)
    document.core_properties.title = "WAMOCON Marketing-Maschine – Betriebs- und Anwenderhandbuch"
    document.core_properties.subject = "Operator- und Administratorhandbuch für die fünf WAMOCON Kampagnen"
    document.core_properties.author = "WAMOCON"
    document.core_properties.last_modified_by = "WAMOCON"
    document.core_properties.keywords = "WAMOCON, Marketing, Betrieb, Kampagnen, n8n, Freigabe"
    document.core_properties.comments = "Kontrollierte Dokumentation; enthält keine Zugangsdaten."

    settings = document.settings._element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)

    _cover(
        document,
        metadata=metadata,
        status_date=status_date,
        status_label=status_label,
    )
    toc_title = document.add_heading("Inhaltsverzeichnis", level=1)
    _set_keep(toc_title, with_next=True)
    toc = document.add_paragraph()
    _add_toc(toc)
    _parse_markdown(document, source_text)

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def main() -> int:
    parser = argparse.ArgumentParser(description="Build the WAMOCON Word handbook from Markdown.")
    parser.add_argument(
        "--source",
        type=Path,
        default=Path("docs/WAMOCON-MARKETING-HANDBOOK.md"),
    )
    parser.add_argument(
        "--output",
        type=Path,
        default=Path("docs/WAMOCON-Marketing-Handbuch.docx"),
    )
    args = parser.parse_args()
    build(args.source, args.output)
    print(args.output.resolve())
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
